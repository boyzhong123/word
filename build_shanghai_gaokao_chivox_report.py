import math
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = "/Users/zhong/Documents/New project/上海高考英语题型与驰声产品设计分析报告.docx"
ASSET_DIR = "/Users/zhong/Documents/New project/outputs/shanghai_gaokao_report_assets"
FONT_PATH = "/System/Library/Fonts/Hiragino Sans GB.ttc"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(95, 108, 125)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"

EAST_ASIA_FONT = "PingFang SC"
LATIN_FONT = "Calibri"
VISUALS = {}


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=LATIN_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=9.0, color=INK, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def set_table_geometry(table, widths_dxa, width_dxa=9360, indent_dxa=120):
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))


def style_table_borders(table, color="C8D2DE"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def paragraph_border_bottom(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=18, after=10, line=1.1)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(p, before=14, after=7, line=1.15)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, before=10, after=5, line=1.15)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=4, line=1.25)
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_note_box(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    style_table_borders(table, color="D6DDE6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.2)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths_dxa, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    style_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=9.2, color=DARK_BLUE, fill=LIGHT_BLUE)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row in rows:
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size, color=INK, fill=WHITE)
    doc.add_paragraph()
    return table


def pil_font(size, bold=False):
    path = "/System/Library/Fonts/STHeiti Medium.ttc" if bold else FONT_PATH
    return ImageFont.truetype(path, size)


def text_size(draw, text, font):
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw, text, font, max_width):
    lines = []
    for para in text.split("\n"):
        line = ""
        for ch in para:
            trial = line + ch
            if draw.textlength(trial, font=font) <= max_width or not line:
                line = trial
            else:
                lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return lines


def draw_wrapped(draw, box, text, font, fill=(11, 37, 69), align="left", line_gap=10):
    x0, y0, x1, y1 = box
    lines = wrap_text(draw, text, font, x1 - x0)
    line_heights = [text_size(draw, line, font)[1] for line in lines]
    total_h = sum(line_heights) + max(0, len(lines) - 1) * line_gap
    y = y0 + max(0, (y1 - y0 - total_h) / 2)
    for line, h in zip(lines, line_heights):
        w, _ = text_size(draw, line, font)
        if align == "center":
            x = x0 + (x1 - x0 - w) / 2
        elif align == "right":
            x = x1 - w
        else:
            x = x0
        draw.text((x, y), line, font=font, fill=fill)
        y += h + line_gap


def rounded_card(draw, xy, fill, outline=(205, 215, 226), radius=24, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, start, end, fill=(46, 116, 181), width=6):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 18
    left = (
        end[0] - arrow_len * math.cos(angle - math.pi / 6),
        end[1] - arrow_len * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - arrow_len * math.cos(angle + math.pi / 6),
        end[1] - arrow_len * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=fill)


def base_canvas(title, subtitle=None, width=1800, height=980):
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), title, font=pil_font(42, bold=True), fill=(31, 77, 120))
    if subtitle:
        draw.text((72, 112), subtitle, font=pil_font(24), fill=(95, 108, 125))
    return img, draw


def save_score_map(path):
    img, draw = base_canvas("上海高考英语150分：题型分值与驰声抓手", "听说是直接主场，读写综合可通过口头输出转化为可测训练。")
    segments = [
        ("听力选择", 25, (46, 116, 181)),
        ("口语表达", 10, (48, 164, 108)),
        ("语言知识", 35, (242, 184, 75)),
        ("阅读理解", 30, (111, 94, 170)),
        ("读写综合", 50, (232, 93, 117)),
    ]
    x, y, bar_w, bar_h = 120, 245, 1560, 145
    cur = x
    for name, score, color in segments:
        w = int(bar_w * score / 150)
        draw.rounded_rectangle((cur, y, cur + w, y + bar_h), radius=20, fill=color)
        draw_wrapped(draw, (cur + 15, y + 24, cur + w - 15, y + bar_h - 24), f"{name}\n{score}分", pil_font(26, bold=True), fill=(255, 255, 255), align="center")
        cur += w
    draw.text((120, 420), "产品机会判断", font=pil_font(30, bold=True), fill=(31, 77, 120))
    cards = [
        ("P0 直接变现", "35分听说：全真模拟、逐词纠音、开放题评分、二次修复。", (230, 246, 240)),
        ("P1 放大价值", "50分读写综合：概要/翻译/作文先口头预演，再进入书面训练。", (244, 246, 249)),
        ("P2 数据沉淀", "全部题型沉淀错因、错音、错句、能力标签，服务教师端。", (244, 246, 249)),
    ]
    card_y = 500
    for i, (head, body, fill) in enumerate(cards):
        x0 = 120 + i * 520
        rounded_card(draw, (x0, card_y, x0 + 470, card_y + 260), fill)
        draw.text((x0 + 30, card_y + 32), head, font=pil_font(28, bold=True), fill=(31, 77, 120))
        draw_wrapped(draw, (x0 + 30, card_y + 92, x0 + 440, card_y + 230), body, pil_font(24), fill=(11, 37, 69), line_gap=12)
    img.save(path)


def save_method_ladder(path):
    img, draw = base_canvas("更强背书体系：用“大众听过 + 教研认可”的方法命名", "把方法名变成产品心智，把研究依据藏在机制里。")
    rows = [
        ("官方课程标准", "英语学习活动观", "学习理解 → 应用实践 → 迁移创新", (232, 238, 245)),
        ("大众学习法", "艾宾浩斯 + 费曼", "间隔复习抗遗忘；讲清楚才是真懂", (246, 251, 248)),
        ("语言习得理论", "克拉申 i+1 + 输出假说", "先给可理解输入，再用输出暴露差距", (255, 248, 232)),
        ("中国英语教研", "文秋芳产出导向法 POA", "输出驱动 → 输入促成 → 评价修复", (244, 246, 249)),
        ("教育评价理论", "布鲁姆掌握学习 + 形成性评价", "达标再进阶；反馈指向下一步", (248, 240, 245)),
        ("高手训练法", "刻意练习", "拆微技能、给即时反馈、反复修正", (237, 247, 252)),
    ]
    left_x, mid_x, right_x = 90, 430, 950
    y0, row_h = 190, 112
    for i, (kind, name, product, fill) in enumerate(rows):
        y = y0 + i * row_h
        rounded_card(draw, (left_x, y, 350, y + 84), fill, radius=18)
        draw_wrapped(draw, (left_x + 20, y + 12, 330, y + 72), kind, pil_font(24, bold=True), fill=(31, 77, 120), align="center")
        rounded_card(draw, (mid_x, y, 850, y + 84), (255, 255, 255), radius=18)
        draw_wrapped(draw, (mid_x + 25, y + 10, 825, y + 74), name, pil_font(26, bold=True), fill=(11, 37, 69), align="center")
        rounded_card(draw, (right_x, y, 1670, y + 84), fill, radius=18)
        draw_wrapped(draw, (right_x + 25, y + 10, 1645, y + 74), product, pil_font(25), fill=(11, 37, 69), align="center")
        draw_arrow(draw, (360, y + 42), (420, y + 42), fill=(120, 140, 160), width=4)
        draw_arrow(draw, (860, y + 42), (940, y + 42), fill=(120, 140, 160), width=4)
    img.save(path)


def save_flywheel(path):
    img, draw = base_canvas("驰声学习闭环飞轮", "让“测评”从一次打分，升级为持续提分系统。")
    center = (900, 520)
    rounded_card(draw, (690, 390, 1110, 650), (232, 238, 245), radius=42, outline=(46, 116, 181), width=5)
    draw_wrapped(draw, (735, 425, 1065, 615), "驰声AI测评\n诊断 + 纠音 + 反馈", pil_font(34, bold=True), fill=(31, 77, 120), align="center", line_gap=16)
    nodes = [
        ("高考题型任务", "按上海题型、时长、分值训练"),
        ("可理解输入 i+1", "难度略高于当前水平"),
        ("输出驱动", "先说/写/译，暴露差距"),
        ("AI多维评分", "发音、流利、内容、语法"),
        ("二次修复", "重读、重答、重写、解释"),
        ("间隔复习", "错题错音自动排程"),
        ("教师画像", "班级短板与作业推荐"),
    ]
    radius_x, radius_y = 660, 310
    points = []
    for i, (title, body) in enumerate(nodes):
        angle = -math.pi / 2 + i * 2 * math.pi / len(nodes)
        px = center[0] + radius_x * math.cos(angle)
        py = center[1] + radius_y * math.sin(angle)
        points.append((px, py))
    for i in range(len(points)):
        sx, sy = points[i]
        ex, ey = points[(i + 1) % len(points)]
        draw_arrow(draw, (sx, sy), (ex, ey), fill=(180, 195, 210), width=4)
    for (px, py), (title, body) in zip(points, nodes):
        x0, y0 = px - 170, py - 74
        rounded_card(draw, (x0, y0, x0 + 340, y0 + 148), (255, 255, 255), radius=24)
        draw_wrapped(draw, (x0 + 20, y0 + 16, x0 + 320, y0 + 58), title, pil_font(23, bold=True), fill=(31, 77, 120), align="center")
        draw_wrapped(draw, (x0 + 24, y0 + 62, x0 + 316, y0 + 126), body, pil_font(19), fill=(11, 37, 69), align="center", line_gap=8)
    img.save(path)


def save_heatmap(path):
    img, draw = base_canvas(
        "题型与教研/学习法机制匹配图",
        "矩阵表示“产品化优先级”：列名采用“强背书姓名/理论名 + 本文简称”，方便和正文表格互相对应。",
        height=1120,
    )
    cols = [
        "王蔷/梅德明\n活动观",
        "克拉申\ni+1输入",
        "文秋芳\nPOA输出",
        "王初明\n以续促学",
        "刻意练习/CF\n纠错反馈",
        "艾宾浩斯/费曼\n间隔+解释",
    ]
    rows = [
        ("听力选择", [2, 3, 1, 1, 2, 2]),
        ("朗读/纠音", [2, 1, 2, 1, 3, 2]),
        ("开放口语", [3, 2, 3, 2, 3, 1]),
        ("语法词汇", [1, 1, 1, 1, 2, 3]),
        ("阅读/六选四", [3, 3, 2, 2, 1, 3]),
        ("概要/翻译/写作", [3, 2, 3, 3, 3, 2]),
    ]
    x0, y0 = 270, 235
    cell_w, cell_h = 225, 98
    draw.text((86, y0 + 26), "题型", font=pil_font(24, bold=True), fill=(31, 77, 120))
    for c, name in enumerate(cols):
        x = x0 + c * cell_w
        rounded_card(draw, (x + 16, y0 - 86, x + cell_w - 22, y0 - 18), (232, 238, 245), outline=(205, 215, 226), radius=18, width=2)
        draw_wrapped(draw, (x + 24, y0 - 80, x + cell_w - 30, y0 - 24), name, pil_font(17, bold=True), fill=(31, 77, 120), align="center", line_gap=5)
    styles = {
        1: ((255, 255, 255), (205, 215, 226), 2, "辅助"),
        2: ((237, 247, 252), (180, 195, 210), 3, "适配"),
        3: ((232, 238, 245), (46, 116, 181), 4, "核心"),
    }
    for r, (row_name, values) in enumerate(rows):
        y = y0 + r * cell_h
        rounded_card(draw, (76, y + 15, 242, y + cell_h - 15), (255, 255, 255), outline=(205, 215, 226), radius=18, width=2)
        draw_wrapped(draw, (88, y + 24, 230, y + cell_h - 24), row_name, pil_font(21, bold=True), fill=(31, 77, 120), align="center")
        for c, val in enumerate(values):
            x = x0 + c * cell_w
            fill, outline, width, label = styles[val]
            draw.rounded_rectangle((x + 20, y + 16, x + cell_w - 30, y + cell_h - 16), radius=20, fill=fill, outline=outline, width=width)
            draw_wrapped(
                draw,
                (x + 30, y + 24, x + cell_w - 40, y + cell_h - 24),
                label,
                pil_font(20, bold=True),
                fill=(11, 37, 69),
                align="center",
            )
    legend_y = y0 + len(rows) * cell_h + 66
    draw.text((250, legend_y + 4), "图例", font=pil_font(23, bold=True), fill=(31, 77, 120))
    for i, level in enumerate([1, 2, 3]):
        x = 380 + i * 310
        fill, outline, width, label = styles[level]
        draw.rounded_rectangle((x, legend_y, x + 92, legend_y + 42), radius=17, fill=fill, outline=outline, width=width)
        draw.text((x + 112, legend_y + 6), label, font=pil_font(21, bold=True), fill=(11, 37, 69))
    draw_wrapped(
        draw,
        (190, 925, 1610, 970),
        "简称对照：活动观=英语学习活动观；POA=产出导向法；TBLT/任务型=任务型语言教学；i+1=略高于当前水平的可理解输入；CF=纠错反馈。",
        pil_font(20),
        fill=(31, 77, 120),
        align="center",
        line_gap=6,
    )
    draw_wrapped(
        draw,
        (145, 1010, 1655, 1080),
        "读写题型不是硬接入口语，而是通过口头解释、复述、口译和提纲，把不可见的理解与表达过程转成可评测数据。",
        pil_font(23),
        fill=(95, 108, 125),
        align="center",
        line_gap=8,
    )
    img.save(path)


def save_summary_bridge(path):
    img, draw = base_canvas("总结落地：从方法背书到产品价值", "把“理论可信”翻译成“功能可做、证据可看、话术可卖”。", height=1080)
    headers = ["背书资产", "产品机制", "可见证据", "对外价值话术"]
    rows = [
        ("新课标活动观\n+ 克拉申i+1", "分级听读输入\n题型任务递进", "i+1完成率\n听读正确率提升", "不是机械刷题\n是核心素养导向训练"),
        ("费曼学习法\n+ POA输出驱动", "先说/解释/复述\n再进入书面任务", "答案理由命中率\n二次修复率", "让不可见思维\n变成可诊断数据"),
        ("艾宾浩斯\n+ 掌握学习", "错题错音复习日历\n微技能达标进阶", "7/14天保持率\n薄弱标签下降", "从一次练会\n变成长久掌握"),
        ("刻意练习\n+ 形成性评价", "拆音素/重音/内容点\n即时反馈再练", "发音项提升曲线\n开放题rubric提升", "从AI打分\n升级到AI提分"),
    ]
    x0, y0 = 90, 210
    col_w = [350, 390, 380, 520]
    row_h = 160
    colors = [(232, 238, 245), (246, 251, 248), (255, 248, 232), (237, 247, 252)]
    x = x0
    for i, h in enumerate(headers):
        rounded_card(draw, (x, y0 - 72, x + col_w[i] - 18, y0 - 10), (232, 238, 245), radius=16)
        draw_wrapped(draw, (x + 12, y0 - 62, x + col_w[i] - 30, y0 - 20), h, pil_font(24, bold=True), fill=(31, 77, 120), align="center")
        x += col_w[i]
    for r, row in enumerate(rows):
        y = y0 + r * row_h
        x = x0
        for c, text in enumerate(row):
            fill = colors[r] if c == 0 else (255, 255, 255)
            outline = (46, 116, 181) if c == 0 else (205, 215, 226)
            rounded_card(draw, (x, y, x + col_w[c] - 18, y + row_h - 24), fill, outline=outline, radius=20, width=3)
            font = pil_font(25, bold=(c == 0 or c == 3))
            draw_wrapped(draw, (x + 18, y + 16, x + col_w[c] - 36, y + row_h - 40), text, font, fill=(11, 37, 69), align="center", line_gap=10)
            if c < 3:
                draw_arrow(draw, (x + col_w[c] - 8, y + row_h / 2 - 12), (x + col_w[c] + 34, y + row_h / 2 - 12), fill=(120, 140, 160), width=4)
            x += col_w[c]
    draw.text((100, 900), "总结判断", font=pil_font(30, bold=True), fill=(31, 77, 120))
    draw_wrapped(
        draw,
        (100, 940, 1680, 1035),
        "驰声最强的抓手不是单次评分，而是把每次听、说、读、写、译都转化为可诊断、可修复、可复习、可被教师看见的学习证据。",
        pil_font(27, bold=True),
        fill=(11, 37, 69),
        align="center",
        line_gap=12,
    )
    img.save(path)


def create_visual_assets():
    os.makedirs(ASSET_DIR, exist_ok=True)
    paths = {
        "score": os.path.join(ASSET_DIR, "score_map.png"),
        "methods": os.path.join(ASSET_DIR, "method_ladder.png"),
        "flywheel": os.path.join(ASSET_DIR, "chivox_flywheel.png"),
        "heatmap": os.path.join(ASSET_DIR, "method_heatmap.png"),
        "summary": os.path.join(ASSET_DIR, "summary_bridge.png"),
    }
    save_score_map(paths["score"])
    save_method_ladder(paths["methods"])
    save_flywheel(paths["flywheel"])
    save_heatmap(paths["heatmap"])
    save_summary_bridge(paths["summary"])
    return paths


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=4, line=1)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, before=0, after=8, line=1.1)
    r = cap.add_run(caption)
    set_run_font(r, size=8.5, color=MUTED, italic=True)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    hdr = section.header.paragraphs[0]
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hdr, before=0, after=0, line=1)
    run = hdr.add_run("上海高考英语题型分析 | 驰声产品设计建议")
    set_run_font(run, size=9, color=MUTED)

    ftr = section.footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(ftr, before=0, after=0, line=1)
    run = ftr.add_run("内部分析简报 | 2026-05-08")
    set_run_font(run, size=9, color=MUTED)
    return doc


def add_cover(doc):
    add_para(doc, "分析报告", size=11, color=MUTED, bold=True, before=8, after=4)
    add_para(doc, "上海高考英语题型与驰声产品设计分析报告", size=22, color=INK, bold=True, before=0, after=4, line=1.15)
    add_para(doc, "基于150分题型结构，映射英语学习方法背书、驰声听说测评能力与可产品化方案", size=12.5, color=MUTED, after=14)

    rows = [
        ("对象", "上海高考英语 / 外语英语方向"),
        ("依据", "用户提供《上海高考英语150分题型、分值及考查重点.docx》；上海市教育考试院/市教委公开文件；驰声官网公开能力口径；学习科学研究文献"),
        ("报告用途", "教研定位、产品规划、销售材料底层论证、上海高考听说/读写融合训练方案设计"),
        ("日期", "2026-05-08"),
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2100, 7260])
    style_table_borders(table, color="D6DDE6")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=9.5, color=DARK_BLUE, fill=LIGHT_GRAY)
        set_cell_text(cells[1], value, size=9.5, color=INK, fill=WHITE)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)


def add_executive_summary(doc):
    add_heading(doc, "一、核心判断", 1)
    add_note_box(
        doc,
        "结论：",
        "上海高考英语的机会点不只在听说35分，而在“听、说、读、写、译、概括”共同指向真实语言运用。驰声最强的口语测评与纠音能力，应从单一听说备考升级为“以输出带动输入、以测评驱动训练、以数据沉淀能力画像”的英语学习系统。",
    )
    add_bullet(doc, "考试结构上，2026官方口径延续外语笔试115分、听说测试35分，总分150分；用户文档中的题型细目与这一总分结构一致。")
    add_bullet(doc, "能力结构上，35分听说直接匹配驰声的考试级口语评测、开放题多维评分、发音诊断纠音；50分读写综合（概要、翻译、大作文）也可通过“口头解释/口头复述/口头作文”转化为可测、可练、可反馈的输出任务。")
    add_bullet(doc, "产品结构上，建议把题型训练拆成三个层次：题型仿真层、学习法引擎层、教师/教研数据层。这样才能让“AI评分”从卖点变成长期学习闭环。")
    add_figure(doc, VISUALS["score"], "图1：上海高考英语150分能力结构与驰声产品抓手")

    add_heading(doc, "二、考试结构与产品优先级", 1)
    rows = [
        ("听力选择", "25分", "短对话、短文/长对话选择", "听前预测、关键词抓取、语篇信息整合", "P0：用听力诊断联动口头复述，形成“听懂-说出-再听”闭环。"),
        ("口语表达", "10分", "朗读、情景提问、看图说话、快速应答、听后回答", "语音语调、流利度、语用反应、内容组织", "P0：驰声核心主场，做全真模拟、逐词纠音、开放题多维反馈。"),
        ("语言知识", "35分", "语法填空、词汇填空、完形", "语法形式、词汇搭配、上下文逻辑", "P1：用检索练习+间隔复习+错因标签，做基础能力补洞。"),
        ("阅读理解", "30分", "ABC篇阅读、六选四", "细节定位、主旨推断、篇章衔接", "P1：阅读后强制口头解释选择理由，提升可诊断性。"),
        ("读写综合", "50分", "概要写作、翻译、指导性写作", "概括、转述、句法转换、观点表达", "P0/P1：用口头摘要、口头翻译、口头作文预演接入驰声表达评估。"),
    ]
    add_matrix(doc, ["板块", "分值", "题型", "考查重点", "产品优先级"], rows, [1300, 900, 2100, 2500, 2560], font_size=8.8)


def add_method_backing(doc):
    add_heading(doc, "三、学习方法背书：从知名方法到产品机制", 1)
    add_figure(doc, VISUALS["methods"], "图2：更适合对外沟通的学习方法背书体系")
    rows = [
        ("新课标活动观", "教育部高中英语课标强调学习理解、应用实践、迁移创新，天然支持“听说读写看”的综合任务设计。", "全部题型", "产品话术从“刷题工具”升级为“核心素养与高考能力统一训练”。"),
        ("艾宾浩斯", "大众熟悉“遗忘曲线”；研究层面用间隔练习证明分散复习优于集中复习。", "词汇、语法、错题、错音", "错题/错音进入智能复习日程，按1天、3天、7天、14天等节奏自动再练。"),
        ("费曼学习法", "“能讲明白才是真懂”的心智很强；学术层面可用自我解释研究做支撑。", "语法、完形、阅读、概要、翻译", "错题后追加口头解释：为什么选、为什么错、如何改；AI判断解释是否命中规则。"),
        ("克拉申 i+1", "语言教师熟悉的可理解输入理论：材料要略高于学生当前水平，但仍可理解。", "听力、阅读、词汇、语篇输入", "按学生水平推送i+1听读材料，先降低理解门槛，再引出输出任务。"),
        ("输出假说/POA", "Swain输出假说与文秋芳产出导向法都强调输出能暴露差距、驱动学习。", "口语、概要、翻译、写作", "把阅读/翻译/作文转成“先说再写”，用驰声测评让输出质量可见。"),
        ("布鲁姆掌握学习", "教育圈熟悉“掌握后再进入下一层级”；适合解释AI个性化学习路径。", "语言知识、朗读、写作", "设置微技能达标线：发音、重音、内容点、句法结构达标后再解锁下一任务。"),
        ("刻意练习", "Ericsson的刻意练习和大众熟悉的“一万小时”概念都强调目标明确、反馈及时。", "朗读、看图说话、翻译、写作", "把总分拆到微技能：音素、重音、停顿、内容点、句法错误；每次只攻一个短板。"),
        ("形成性评价", "Black、Wiliam、Hattie等反馈研究适合支撑“测评即学习”的产品定位。", "所有题型", "报告不只给分，要给目标、当前状态和下一步练法，形成评学一体。"),
    ]
    add_matrix(doc, ["强背书名称", "客户可理解的背书要点", "适配题型", "产品化机制"], rows, [1400, 3300, 1900, 2760], font_size=8.45)
    add_figure(doc, VISUALS["heatmap"], "图3：不同题型与教研/学习法机制的产品化匹配强度")


def add_question_type_matrices(doc):
    doc.add_page_break()
    add_heading(doc, "四、逐题型映射：考什么、怎么学、驰声怎么做", 1)
    add_figure(doc, VISUALS["question_product_map"], "图4：上海高考英语题型到学生App练法与驰声能力的产品化总览")

    add_heading(doc, "A. 听说部分（35分）", 2)
    rows = [
        ("听短对话选择\n10分", "捕捉身份、地点、态度、时间、数字与隐含意图。", "克拉申i+1：分级可理解输入；艾宾浩斯：错题复听间隔复现。", "听力选择后增加口头复述/解释，系统比对关键词，诊断“没听到、听错、推断错”。"),
        ("听短文或长对话选择\n15分", "在较长语篇中整合信息，理解因果、转折、说话人态度。", "新课标活动观：从理解到应用；克拉申i+1：长音频分层输入。", "长音频切片训练；自动生成听力画像：细节、主旨、推断、数字、态度。"),
        ("朗读句子\n1分", "单词发音、重音、连读、语调、完整度。", "刻意练习：拆到音素/重音；形成性评价：即时反馈后复读。", "使用音素级/逐词诊断纠音，给出漏读、多读、错读、重音与语调反馈。"),
        ("朗读段落\n1分", "段落流利度、停顿、节奏、语义群与稳定输出。", "刻意练习+掌握学习：意群、停顿、语速达标后进阶。", "可视化语速、停顿、完整度和句群节奏，生成段落朗读复测任务。"),
        ("情景提问\n2分", "根据情境生成合适问题，考查疑问句结构和交际目的。", "输出假说/POA：先产出暴露差距；费曼：解释为什么这样问。", "情景问答引擎+语法/内容评分，给出可替换句型和二次追问训练。"),
        ("看图说话\n1.5分", "观察信息、组织顺序、连贯描述、合理推断。", "新课标迁移创新+POA：真实情境驱动表达。", "开放题多维评分：内容、发音、语法、流利度；图片题库自动生成内容点清单。"),
        ("快速应答\n2分", "听懂刺激句并在短时限内自然回应。", "艾宾浩斯+刻意练习：高频功能句块间隔复现、限时自动化。", "5秒限时闯关，记录反应时、完整度和语用是否匹配，形成自动化程度指标。"),
        ("听短文回答问题\n2.5分", "听力理解与口头表达结合，回答要准确、简洁、完整。", "克拉申输入+输出假说：听懂后必须说出，差距进入反馈。", "音频理解+口头答案内容评分；把漏答信息点推送到复听片段。"),
    ]
    add_matrix(doc, ["题型", "考查重点", "学习方法背书", "驰声结合点/产品设计"], rows, [1250, 2500, 2600, 3010], font_size=8.35)

    add_heading(doc, "B. 笔试部分（115分）", 2)
    rows = [
        ("语法填空\n10分", "在语境中判断时态、语态、非谓语、从句、词形。", "费曼学习法：说清规则；艾宾浩斯：同错因间隔复现。", "错题后让学生口头解释规则，AI检查解释关键词；生成同规则变式题。"),
        ("词汇填空\n10分", "词义辨析、词性、搭配、上下文语义场。", "艾宾浩斯+克拉申i+1：词块在可理解语篇中复现。", "词汇不只背释义，要求口头造句并评分发音/语法；错词进入复习日程。"),
        ("完形填空\n15分", "上下文逻辑、词汇搭配、语篇连贯、情感态度。", "费曼学习法：解释上下文线索；掌握学习：按错因达标。", "将选项错因标为词义、搭配、逻辑、指代；让学生口头解释上下文线索。"),
        ("阅读ABC篇\n22分", "细节定位、主旨概括、推断判断、作者态度。", "新课标活动观+费曼：从理解到口头说明答案理由。", "阅读后1分钟口头摘要+答案理由说明，补足纯选择题不可观测的思维过程。"),
        ("六选四\n8分", "篇章结构、衔接、指代、转折、例证与段落功能。", "费曼+新课标迁移：说清前后线索和段落功能。", "拖拽排序/填句后要求口头说明“前后线索”，AI识别连接词、指代词、主题一致性。"),
        ("概要写作\n10分", "提炼主旨和要点，用60词内转述，不堆细节。", "POA输出驱动+费曼：先讲给别人听，再压缩成文。", "先口头摘要再书面摘要；评分内容覆盖、冗余、字数、语言准确性。"),
        ("中译英\n15分", "句法转换、固定句型、逻辑关系、表达准确。", "刻意练习+布鲁姆掌握：高频句型达标后进阶。", "中译英可先口头翻译，系统检测语法、发音与表达，再进入书面版本。"),
        ("指导性写作\n25分", "审题、立意、结构、论证、语言质量和文化沟通。", "POA+形成性评价：输出、反馈、修订循环。", "写前口头提纲/口头作文，系统评内容、逻辑、语法、流利度；写后给rubric反馈与改写任务。"),
    ]
    add_matrix(doc, ["题型", "考查重点", "学习方法背书", "驰声结合点/产品设计"], rows, [1250, 2500, 2600, 3010], font_size=8.35)


def add_product_design(doc):
    add_heading(doc, "五、驰声产品设计：把“AI测评”变成学习闭环", 1)
    add_note_box(
        doc,
        "产品原则：",
        "所有题型都要回答三个问题：学生现在会什么、为什么丢分、下一次怎么练。驰声的竞争力不应只呈现为“能打分”，而应呈现为“能诊断、能纠错、能安排复习、能让老师看见进步”。",
    )
    add_figure(doc, VISUALS["flywheel"], "图4：驰声从测评工具升级为学习闭环的产品逻辑")

    rows = [
        ("题型仿真层", "还原上海听说/笔试题型、时长、分值和答题流程。", "全真模考、分题型专项、限时训练、官方口径分值看板。", "让学校和学生相信产品“对考试”。"),
        ("学习法引擎层", "把知名方法背书做成产品动作，而非文案。", "艾宾浩斯复习、费曼解释、克拉申i+1输入、POA输出驱动、布鲁姆达标。", "让训练真正提分，减少刷题疲劳。"),
        ("AI测评反馈层", "把总分拆成可行动指标。", "发音、完整度、流利度、语调、内容、语法、逻辑、反应时。", "形成驰声区别于普通题库的核心体验。"),
        ("教师/教研层", "把学生数据汇成班级教学决策。", "错因分布、能力画像、推荐作业、区域/班级测评报告。", "服务校内采购、区域测评与教研闭环。"),
    ]
    add_matrix(doc, ["层级", "定位", "关键功能", "商业价值"], rows, [1600, 2550, 2950, 2260], font_size=8.8)

    add_heading(doc, "建议产品包", 2)
    add_bullet(doc, "上海高考听说冲刺舱：按35分结构提供全真模拟、逐项评分、错音/错因复练、二次提交，主打驰声考试级评测和开放题多维评分。")
    add_bullet(doc, "读写口融合训练营：阅读后口头复述，概要先说后写，翻译先说后写，作文先口头提纲再成文，用口语测评把不可见思维外显。")
    add_bullet(doc, "AI错因画像与复习日历：错题、错词、错音、错句自动进入间隔复习；教师端看到班级共性短板。")
    add_bullet(doc, "题型能力图谱：每道题绑定微技能标签，例如“态度推断、非谓语、指代衔接、内容点覆盖、重音、停顿、反应时”。")
    add_bullet(doc, "二次修复机制：每次反馈后必须重答、复读、重写或解释，系统记录修复率，作为比单次分数更能体现学习成长的指标。")

    add_heading(doc, "六、路线图与验证指标", 1)
    rows = [
        ("P0 0-3个月", "上海高考听说35分题型仿真；开放题多维评分；逐词纠音；听后回答内容评分。", "模考完成率、二次提交率、朗读/开放题提分、学生周活。"),
        ("P1 3-6个月", "读写口融合：阅读口头摘要、概要写作评分、翻译口头预演、写作rubric反馈。", "概要要点覆盖率、翻译结构正确率、写作二次修改提升幅度。"),
        ("P2 6-9个月", "学习法引擎：间隔复习、检索练习、自我解释评分、交错练习推荐。", "错题复现正确率、7/14天保持率、薄弱标签下降率。"),
        ("P3 9-12个月", "学校/区域教研：班级报告、校本题库、区域模考、教师推荐作业。", "教师布置率、班级报告使用率、续费/扩校线索、区域测评覆盖。"),
    ]
    add_matrix(doc, ["阶段", "产品重点", "验证指标"], rows, [1600, 4650, 3110], font_size=8.8)
    add_figure(doc, VISUALS["phase_metrics"], "图9：P0-P3路线图与验证指标看板")


def add_conclusion(doc):
    add_heading(doc, "七、总结：把方法背书翻译成产品价值", 1)
    add_note_box(
        doc,
        "一句话定位：",
        "驰声可以把上海高考英语产品定位为“官方题型仿真 + 知名学习法引擎 + AI口语测评反馈”的综合提分系统，而不是单一听说模考工具。",
    )
    add_figure(doc, VISUALS["summary"], "图5：方法背书转化为产品机制、数据证据与对外话术")
    add_heading(doc, "落地链路", 2)
    rows = [
        ("背书入口", "对外先讲新课标、艾宾浩斯、费曼、克拉申i+1、POA、布鲁姆、刻意练习、形成性评价，建立可信心智。"),
        ("功能承接", "每个背书必须落到一个功能：分级输入、先说后写、错因解释、复习日历、达标进阶、二次修复、教师报告。"),
        ("数据证明", "每个功能必须有一个证据指标：完成率、二次提交率、错因复现正确率、7/14天保持率、rubric提升、班级短板下降。"),
        ("销售表达", "不要只说“AI评分准”，要说“AI把学生为什么错、怎么改、是否真的掌握，变成老师和学校看得见的数据”。"),
    ]
    add_matrix(doc, ["层级", "具体落地"], rows, [1450, 7910], font_size=8.9)
    add_heading(doc, "四方价值", 2)
    rows = [
        ("对学生", "从“多刷题”变成“知道该听什么、说什么、为什么错、下次怎么改”，降低无效努力。"),
        ("对老师", "把错因、错音、错句和班级共性问题自动沉淀出来，减少批改负担，提升布置作业的精准度。"),
        ("对学校", "获得能服务考试、课堂和区域测评的数据化英语学习闭环，而不是只购买一套题库。"),
        ("对驰声", "核心竞争力从“口语评分准确”升级为“用口语测评外显英语学习过程，并驱动全题型提分”。"),
    ]
    add_matrix(doc, ["对象", "最终价值"], rows, [1600, 7760], font_size=9.0)
    add_bullet(doc, "背书层建议对外主推：新课标英语学习活动观、艾宾浩斯、费曼学习法、克拉申i+1、输出假说/产出导向法、布鲁姆掌握学习、刻意练习、形成性评价。")
    add_bullet(doc, "产品层建议少讲“我们有AI评分”，多讲“每一道题都有诊断、修复、复习和教师可见的数据证据”。")
    add_bullet(doc, "销售层建议把35分听说作为入口，把50分读写综合的口头预演作为差异化增量，把教师端报告作为校内采购理由。")


def add_learning_method_appendix(doc):
    doc.add_page_break()
    add_heading(doc, "八、附录A：学习法背书卡片", 1)
    add_note_box(
        doc,
        "使用建议：",
        "对外沟通时优先使用“官方/权威解释 + 产品机制 + 数据证据”的三段式表达。费曼学习法和艾宾浩斯适合做大众心智入口；新课标活动观、POA、克拉申i+1、布鲁姆掌握学习和形成性评价更适合做教研背书。",
    )
    rows = [
        (
            "新课标英语学习活动观",
            "高中英语教学的官方课程理念，强调学生在主题意义引领下，通过活动学习和运用语言。",
            "教育部《普通高中英语课程标准》提出指向核心素养发展的英语学习活动观，核心活动包括学习理解、应用实践、迁移创新。",
            "这是学校和教研最容易接受的“官方背书”，可解释为什么产品不能只刷选择题，而要训练听说读写看综合任务。",
            "题型仿真、读后口头复述、概要先说后写、情境问答、看图说话、教师报告。",
        ),
        (
            "艾宾浩斯遗忘曲线",
            "大众熟悉的记忆规律：新学内容如果不被再次激活，会快速遗忘。",
            "Ebbinghaus在1885年对记忆和遗忘做系统实验；现代间隔练习研究进一步支持分散复习优于集中复习。",
            "它能为“错题/错音不能只练一次”背书，适合转化为智能复习日历和学习保持率指标。",
            "错题、错词、错音、错句自动排程；按1天、3天、7天、14天复练；追踪7/14天保持率。",
        ),
        (
            "费曼学习法",
            "用简单语言把知识讲给别人听，讲不清楚的地方就是理解漏洞。",
            "大学学习中心通常把它解释为：选择概念、用简单语言解释、发现理解缺口、回到材料补足并再次简化。严格说它不是官方课程方法，学术支撑可用自我解释研究补强。",
            "它很适合做大众传播，也适合解释“为什么选择题做完后还要口头说理由”。",
            "阅读答案理由说明、语法错题口头解释、完形上下文线索解释、概要口头讲给别人听。",
        ),
        (
            "克拉申 i+1 可理解输入",
            "语言输入要比学生当前水平略难一点，但仍能理解；太简单无增长，太难会放弃。",
            "Krashen输入假说常被概括为i+1：i是学习者当前水平，+1是略高一级的语言输入。",
            "它能支撑分层听力/阅读材料和个性化难度推荐，是“不是所有学生做同一套材料”的理论依据。",
            "按能力推送i+1听读材料；长音频切片；词汇在可理解语篇中复现；阅读后输出任务。",
        ),
        (
            "Swain输出假说",
            "输入不够，学生必须尝试表达，才能发现自己说不出、说不准、说不完整的地方。",
            "Swain在二语习得研究中提出输出在交际能力发展中有重要作用，输出能促进注意语言差距和修正表达。",
            "它直接支撑驰声优势：口语测评不是附加题，而是让学习差距显性的关键入口。",
            "听后回答、看图说话、口头摘要、口头翻译、写前口头提纲、二次重答。",
        ),
        (
            "文秋芳产出导向法 POA",
            "中国外语教育领域较有影响力的本土理论，主张用产出任务驱动学习，再用输入材料促成产出。",
            "POA通常被概括为输出驱动、输入促成、评价；相关研究强调驱动、促成、评价的循环链。",
            "相比单独讲Swain，POA更适合中国学校场景，能解释“先说后写、先产出再促成”的产品结构。",
            "读写口融合训练营；阅读先复述、概要先口头摘要、翻译先口译、作文先口头提纲再成文。",
        ),
        (
            "布鲁姆掌握学习",
            "学生在前置知识或微技能掌握前，不应直接跳到下一层难度。",
            "Bloom的Learning for Mastery强调学习需要合适的教学质量、理解任务、坚持和足够时间，学习进阶应围绕掌握情况调整。",
            "它能为AI个性化路径背书：不是按天数推进，而是按微技能达标推进。",
            "音素/重音/停顿/内容点/句法结构达标线；达标解锁下一题型；薄弱标签补练。",
        ),
        (
            "刻意练习",
            "把大目标拆成小技能，在明确目标、即时反馈和重复修正中提升表现。",
            "Ericsson等关于专家表现的研究强调，持续进步来自目标明确、反馈及时、难度适中的练习。",
            "它能支撑驰声把总分拆成微技能：发音、重音、停顿、完整度、内容点、语法。",
            "朗读逐词纠音、看图说话内容点覆盖、翻译高频句型、作文结构和表达专项训练。",
        ),
        (
            "形成性评价",
            "评价不是学习结束后的判分，而是学习过程中的反馈和下一步行动。",
            "Black与Wiliam的Assessment for Learning、Hattie与Timperley的反馈研究都强调反馈应告诉学生目标、当前状态和下一步。",
            "它是“AI测评即学习”的核心背书：分数只是起点，反馈和修复才是价值。",
            "每题报告包含目标、当前表现、错因、下一步练法；教师端班级短板和推荐作业。",
        ),
        (
            "口语纠错反馈",
            "二语课堂中，教师或系统对发音、语法、语义、语用错误给出纠正，并促使学习者修复。",
            "Lyster、Saito、Sato等综述指出口语纠错反馈需要结合错误类型、学习者回应和具体语境。",
            "这与驰声口语测评最贴近：纠音不能停在标红，要引导学生重读、重答、修复。",
            "逐词纠音、漏读/多读/错读反馈、开放口语内容纠错、二次提交、修复率追踪。",
        ),
    ]
    add_matrix(
        doc,
        ["学习法", "简要介绍", "官方/权威解释", "背书原因", "驰声产品落点"],
        rows,
        [1200, 1850, 2600, 1800, 1910],
        font_size=7.45,
    )


def add_sources(doc):
    add_heading(doc, "九、来源与引用", 1)
    sources = [
        ("上海市教育考试院：2026春考实施办法，外语115+35。", "https://www.shmeea.edu.cn/page/08000/20251118/19863.html"),
        ("上海市教育考试院：2026春考问答。", "https://www.shmeea.edu.cn/page/02300/20251118/19859.html"),
        ("上海市教委：2025春考招生试点方案。", "https://www.shanghai.gov.cn/cmsres/a2/a2a03877be6c4e9a87660c610d19af05/2c1edede8fa07e170792567a9bf489d6.pdf"),
        ("用户提供题型文档。", "/Users/zhong/Downloads/上海高考英语150分题型、分值及考查重点.docx"),
        ("教育部：普通高中英语课程标准。", "https://www.pep.com.cn/xw/zt/rjwy/gzkb2020/202205/P020220517522153664167.pdf"),
        ("驰声官网：公司简介与考试级评测引擎。", "https://www.chivox.com/gettoknowus/index.aspx"),
        ("驰声官网：语音评测技术。", "https://www.chivox.com/voiceevaluation/index.aspx"),
        ("驰声官网：开放题多维评估。", "https://www.chivox.com/multidimensionalevaluationofopenquestiontype/index.aspx"),
        ("驰声官网：听说在线闭环。", "https://www.chivox.com/soundheardonline/index.aspx"),
        ("Cepeda et al. (2006)：间隔练习综述。", "https://cir.nii.ac.jp/crid/1361418520534596992"),
        ("Chi et al. (1994)：自我解释。", "https://cir.nii.ac.jp/crid/1363388846185476864"),
        ("Krashen输入假说/i+1。", "https://www.cambridge.org/core/journals/journal-of-classics-teaching/article/comprehensible-input-and-krashens-theory/2308987050E8D31E3986B530D4B02F6F"),
        ("Swain输出假说。", "https://www.semanticscholar.org/paper/Communicative-competence-%3A-Some-roles-of-input-and-Swain/fb4d094b6a03f2ae90e007d4117bd62224702317"),
        ("文秋芳产出导向法POA。", "https://www.atlantis-press.com/proceedings/ichess-21/125967210"),
        ("Bloom掌握学习。", "https://eric.ed.gov/?id=eD053419"),
        ("Ericsson et al. (1993)：刻意练习。", "https://cir.nii.ac.jp/crid/1363670319559953152"),
        ("Black & Wiliam：形成性评价。", "https://www.nuffieldfoundation.org/project/the-assessment-reform-group"),
        ("Lyster et al. (2013)：口语纠错反馈。", "https://www.cambridge.org/core/journals/language-teaching/article/oral-corrective-feedback-in-second-language-classrooms/B33FC71C12A1317DCA0CE704B7B0BC00"),
    ]
    split_at = math.ceil(len(sources) / 2)
    left = sources[:split_at]
    right = sources[split_at:]
    row_count = max(len(left), len(right))
    table = doc.add_table(rows=row_count, cols=2)
    set_table_geometry(table, [4680, 4680])
    for row_idx in range(row_count):
        for col_idx, item in enumerate((left[row_idx], right[row_idx] if row_idx < len(right) else None)):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, top=35, start=80, bottom=35, end=80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.04)
            if item is None:
                continue
            source_idx = row_idx + 1 if col_idx == 0 else row_idx + split_at + 1
            label, url = item
            r = p.add_run(f"[{source_idx}] {label} ")
            set_run_font(r, size=7.4, color=INK)
            if url.startswith("http"):
                add_hyperlink(p, "链接", url)
            else:
                r = p.add_run(url)
                set_run_font(r, size=7.4, color=MUTED)


def build():
    global VISUALS
    VISUALS = create_visual_assets()
    doc = setup_doc()
    add_cover(doc)
    add_executive_summary(doc)
    add_method_backing(doc)
    add_question_type_matrices(doc)
    add_product_design(doc)
    add_conclusion(doc)
    add_sources(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()